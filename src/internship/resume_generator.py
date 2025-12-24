"""
Resume Generator for JARVIS Internship Automation Module.

Generates customized resumes for specific job applications:
- Analyzes job requirements
- Retrieves relevant content via RAG
- Generates tailored sections
- Optimizes for ATS
- Outputs multiple formats
"""

import json
import os
from dataclasses import dataclass, field
from datetime import datetime
from enum import Enum
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

from loguru import logger

from .models import (
    InternshipListing,
    MasterResume,
    GeneratedResume,
    Project,
    Skill,
    UserProfile,
)
from .resume_rag import ResumeRAG, ResumeRAGContext
from .prompts import InternshipPrompts


class ResumeFormat(Enum):
    """Output format for resume."""
    PDF = "pdf"
    DOCX = "docx"
    MARKDOWN = "markdown"
    PLAIN_TEXT = "txt"
    JSON = "json"


@dataclass
class ResumeGenerationConfig:
    """Configuration for resume generation."""
    include_summary: bool = True
    include_education: bool = True
    include_experience: bool = True
    include_projects: bool = True
    include_skills: bool = True
    include_certifications: bool = False
    
    max_projects: int = 4
    max_experience: int = 3
    max_skills_per_category: int = 8
    
    ats_optimization: bool = True
    target_ats_score: float = 85.0
    
    output_formats: List[ResumeFormat] = field(default_factory=lambda: [
        ResumeFormat.PDF,
        ResumeFormat.DOCX,
        ResumeFormat.MARKDOWN,
    ])
    
    output_directory: str = "data/generated_resumes"


class ResumeGenerator:
    """
    Generate customized resumes for job applications.
    
    Workflow:
    1. ANALYZE - Parse job posting, extract requirements
    2. MATCH - Query RAG for matching content
    3. CUSTOMIZE - Generate tailored sections
    4. OPTIMIZE - ATS keyword optimization
    5. OUTPUT - Generate multiple formats
    """
    
    def __init__(
        self,
        profile: Optional[UserProfile] = None,
        master_resume: Optional[MasterResume] = None,
        resume_rag: Optional[ResumeRAG] = None,
        llm_router=None,
        config: Optional[ResumeGenerationConfig] = None,
    ):
        self.profile = profile or UserProfile()
        self.master_resume = master_resume
        self.rag = resume_rag or ResumeRAG()
        self.llm = llm_router
        self.config = config or ResumeGenerationConfig()
        
        # Ensure output directory exists
        os.makedirs(self.config.output_directory, exist_ok=True)
    
    # =========================================================================
    # Main Generation Workflow
    # =========================================================================
    
    async def generate_for_job(
        self,
        job: InternshipListing,
        custom_instructions: Optional[str] = None,
    ) -> GeneratedResume:
        """
        Generate a customized resume for a specific job.
        
        Args:
            job: The internship listing to target
            custom_instructions: Any additional customization instructions
            
        Returns:
            GeneratedResume with all generated content and files
        """
        logger.info(f"Generating resume for {job.company} - {job.role}")
        
        # Step 1: Analyze job posting
        job_analysis = await self._analyze_job(job)
        
        # Step 2: Retrieve matching content
        rag_context = self.rag.retrieve_for_job(
            job_description=job.description,
            job_requirements=job.requirements,
            n_projects=self.config.max_projects,
            n_experience=self.config.max_experience,
        )
        
        # Step 3: Generate customized sections
        sections = await self._generate_sections(job, job_analysis, rag_context)
        
        # Step 4: ATS optimization
        if self.config.ats_optimization:
            sections, ats_score, keywords = await self._optimize_ats(
                sections, job_analysis
            )
        else:
            ats_score = 0.0
            keywords = {"included": [], "missing": []}
        
        # Step 5: Create GeneratedResume object
        generated = GeneratedResume(
            company=job.company,
            role=job.role,
            internship_id=job.id,
            content=sections,
            ats_score=ats_score,
            keywords_included=keywords.get("included", []),
            keywords_missing=keywords.get("missing", []),
            projects_used=[p.id for p, _ in rag_context.matched_projects],
        )
        
        # Step 6: Generate output files
        generated = await self._generate_outputs(generated, job)
        
        logger.info(f"Resume generated with ATS score: {ats_score:.0f}%")
        
        return generated
    
    # =========================================================================
    # Step 1: Job Analysis
    # =========================================================================
    
    async def _analyze_job(self, job: InternshipListing) -> Dict[str, Any]:
        """Analyze job posting to extract requirements and keywords."""
        analysis = {
            "company": job.company,
            "role": job.role,
            "required_skills": [],
            "preferred_skills": [],
            "keywords": job.keywords.copy() if job.keywords else [],
            "responsibilities": job.responsibilities.copy() if job.responsibilities else [],
        }
        
        # Extract keywords from description
        if job.description:
            extracted = self._extract_keywords_from_text(job.description)
            analysis["keywords"].extend(extracted)
        
        # Extract from requirements
        for req in job.requirements:
            extracted = self._extract_keywords_from_text(req)
            analysis["required_skills"].extend(extracted)
        
        # Use LLM for deeper analysis if available
        if self.llm and job.description:
            try:
                prompt = InternshipPrompts.format_template(
                    "analyze_job",
                    job_posting=job.description[:2000]
                )
                
                response = await self.llm.generate(prompt)
                
                # Try to parse JSON response
                try:
                    llm_analysis = json.loads(response)
                    analysis["required_skills"].extend(llm_analysis.get("required_skills", []))
                    analysis["preferred_skills"].extend(llm_analysis.get("preferred_skills", []))
                    analysis["keywords"].extend(llm_analysis.get("keywords", []))
                except json.JSONDecodeError:
                    pass
                    
            except Exception as e:
                logger.warning(f"LLM job analysis failed: {e}")
        
        # Deduplicate
        analysis["required_skills"] = list(set(analysis["required_skills"]))
        analysis["preferred_skills"] = list(set(analysis["preferred_skills"]))
        analysis["keywords"] = list(set(analysis["keywords"]))
        
        return analysis
    
    def _extract_keywords_from_text(self, text: str) -> List[str]:
        """Extract technical keywords from text."""
        keywords = []
        text_lower = text.lower()
        
        # Technical skills to look for
        tech_terms = [
            "python", "java", "javascript", "typescript", "c++", "c#", "go", "rust", "r",
            "sql", "nosql", "mongodb", "postgresql", "mysql", "redis",
            "machine learning", "ml", "deep learning", "ai", "artificial intelligence",
            "data science", "data analysis", "analytics", "statistics", "statistical",
            "tensorflow", "pytorch", "keras", "scikit-learn", "sklearn",
            "pandas", "numpy", "scipy", "matplotlib", "seaborn",
            "spark", "hadoop", "kafka", "airflow", "etl",
            "aws", "gcp", "azure", "cloud",
            "docker", "kubernetes", "ci/cd", "jenkins",
            "git", "github", "gitlab", "version control",
            "react", "angular", "vue", "node.js", "express",
            "api", "rest", "graphql", "microservices",
            "agile", "scrum", "jira",
            "nlp", "natural language", "computer vision",
            "regression", "classification", "clustering", "neural network",
            "linux", "unix", "bash",
        ]
        
        for term in tech_terms:
            if term in text_lower:
                keywords.append(term)
        
        return keywords
    
    # =========================================================================
    # Step 2-3: Generate Sections
    # =========================================================================
    
    async def _generate_sections(
        self,
        job: InternshipListing,
        analysis: Dict[str, Any],
        context: ResumeRAGContext,
    ) -> Dict[str, Any]:
        """Generate all resume sections."""
        sections = {}
        
        # Contact section (from profile)
        sections["contact"] = {
            "name": self.profile.name,
            "email": self.profile.email,
            "phone": self.profile.phone,
            "linkedin": self.profile.linkedin,
            "github": self.profile.github,
            "portfolio": self.profile.portfolio,
        }
        
        # Summary section
        if self.config.include_summary:
            sections["summary"] = await self._generate_summary(job, analysis, context)
        
        # Education section
        if self.config.include_education:
            sections["education"] = self._generate_education()
        
        # Experience section
        if self.config.include_experience and context.matched_experience:
            sections["experience"] = await self._generate_experience(
                job, analysis, context
            )
        
        # Projects section
        if self.config.include_projects and context.matched_projects:
            sections["projects"] = await self._generate_projects(
                job, analysis, context
            )
        
        # Skills section
        if self.config.include_skills:
            sections["skills"] = self._generate_skills(analysis, context)
        
        return sections
    
    async def _generate_summary(
        self,
        job: InternshipListing,
        analysis: Dict[str, Any],
        context: ResumeRAGContext,
    ) -> str:
        """Generate professional summary."""
        if self.llm:
            try:
                prompt = InternshipPrompts.format_template(
                    "resume_summary",
                    name=self.profile.name,
                    university=self.profile.university,
                    major=self.profile.major,
                    year=self.profile.year,
                    gpa=self.profile.gpa,
                    skills=", ".join(self.profile.primary_skills),
                    company=job.company,
                    role=job.role,
                    requirements=", ".join(analysis.get("required_skills", [])[:5]),
                )
                
                return await self.llm.generate(prompt)
                
            except Exception as e:
                logger.warning(f"LLM summary generation failed: {e}")
        
        # Fallback to template
        return (
            f"Data Science student at {self.profile.university} with expertise in "
            f"{', '.join(self.profile.primary_skills[:3])}. "
            f"Seeking {job.role} position to apply analytical skills and contribute to impactful projects."
        )
    
    def _generate_education(self) -> List[Dict[str, Any]]:
        """Generate education section."""
        return [{
            "institution": self.profile.university,
            "degree": "Bachelor of Arts",
            "major": self.profile.major,
            "graduation": f"Expected {self.profile.graduation_year}",
            "gpa": f"{self.profile.gpa:.2f}",
            "relevant_coursework": [
                "Machine Learning",
                "Data Structures",
                "Statistical Methods",
                "Linear Algebra",
            ],
        }]
    
    async def _generate_experience(
        self,
        job: InternshipListing,
        analysis: Dict[str, Any],
        context: ResumeRAGContext,
    ) -> List[Dict[str, Any]]:
        """Generate experience section with tailored bullets."""
        experiences = []
        
        for exp, score in context.matched_experience[:self.config.max_experience]:
            exp_dict = {
                "company": exp.company,
                "role": exp.role,
                "location": exp.location,
                "dates": f"{exp.start_date} - {'Present' if exp.is_current else exp.end_date}",
                "bullets": exp.achievements[:4],
                "match_score": score,
            }
            experiences.append(exp_dict)
        
        return experiences
    
    async def _generate_projects(
        self,
        job: InternshipListing,
        analysis: Dict[str, Any],
        context: ResumeRAGContext,
    ) -> List[Dict[str, Any]]:
        """Generate projects section."""
        projects = []
        
        for project, score in context.matched_projects[:self.config.max_projects]:
            proj_dict = {
                "name": project.name,
                "technologies": project.technologies[:5],
                "bullets": project.resume_bullets[:3] if project.resume_bullets else [project.description],
                "github": project.github_url,
                "match_score": score,
            }
            projects.append(proj_dict)
        
        # If LLM available, enhance project bullets
        if self.llm and projects:
            try:
                # Format projects for prompt
                projects_text = "\n\n".join([
                    f"**{p['name']}**\nTechnologies: {', '.join(p['technologies'])}\n"
                    f"Description: {' '.join(p['bullets'])}"
                    for p in projects
                ])
                
                prompt = InternshipPrompts.format_template(
                    "resume_projects",
                    company=job.company,
                    role=job.role,
                    requirements=", ".join(analysis.get("required_skills", [])[:5]),
                    keywords=", ".join(analysis.get("keywords", [])[:10]),
                    projects=projects_text,
                )
                
                enhanced = await self.llm.generate(prompt)
                # Parse enhanced projects (simplified - in production would parse properly)
                
            except Exception as e:
                logger.warning(f"LLM project enhancement failed: {e}")
        
        return projects
    
    def _generate_skills(
        self,
        analysis: Dict[str, Any],
        context: ResumeRAGContext,
    ) -> Dict[str, List[str]]:
        """Generate skills section organized by category."""
        skills = {
            "Programming Languages": [],
            "Data Science & ML": [],
            "Tools & Frameworks": [],
            "Databases": [],
            "Other": [],
        }
        
        # Categorize skills
        categorization = {
            "Programming Languages": ["python", "java", "javascript", "c++", "r", "sql", "typescript"],
            "Data Science & ML": ["machine learning", "deep learning", "tensorflow", "pytorch", 
                                  "scikit-learn", "pandas", "numpy", "statistics", "nlp"],
            "Tools & Frameworks": ["git", "docker", "kubernetes", "aws", "gcp", "azure", 
                                   "react", "node.js", "spark"],
            "Databases": ["postgresql", "mysql", "mongodb", "redis", "nosql"],
        }
        
        # Add matched skills first (prioritize job requirements)
        all_keywords = set(analysis.get("required_skills", []) + analysis.get("keywords", []))
        
        for skill in context.matched_skills:
            skill_lower = skill.name.lower()
            categorized = False
            
            for category, terms in categorization.items():
                if any(term in skill_lower for term in terms):
                    if skill.name not in skills[category]:
                        skills[category].append(skill.name)
                    categorized = True
                    break
            
            if not categorized:
                if skill.name not in skills["Other"]:
                    skills["Other"].append(skill.name)
        
        # Add profile skills
        for skill in self.profile.primary_skills:
            skill_lower = skill.lower()
            for category, terms in categorization.items():
                if any(term in skill_lower for term in terms):
                    if skill not in skills[category]:
                        skills[category].append(skill)
                    break
        
        # Limit per category
        for category in skills:
            skills[category] = skills[category][:self.config.max_skills_per_category]
        
        # Remove empty categories
        skills = {k: v for k, v in skills.items() if v}
        
        return skills
    
    # =========================================================================
    # Step 4: ATS Optimization
    # =========================================================================
    
    async def _optimize_ats(
        self,
        sections: Dict[str, Any],
        analysis: Dict[str, Any],
    ) -> Tuple[Dict[str, Any], float, Dict[str, List[str]]]:
        """Optimize resume for ATS systems."""
        # Get all text from resume
        resume_text = self._sections_to_text(sections).lower()
        
        # Check keyword coverage
        all_keywords = set(
            analysis.get("required_skills", []) +
            analysis.get("preferred_skills", []) +
            analysis.get("keywords", [])
        )
        
        included = []
        missing = []
        
        for keyword in all_keywords:
            if keyword.lower() in resume_text:
                included.append(keyword)
            else:
                missing.append(keyword)
        
        # Calculate score
        if all_keywords:
            coverage = len(included) / len(all_keywords)
            ats_score = coverage * 100
        else:
            ats_score = 80.0  # Default if no keywords
        
        # Use LLM for optimization suggestions if score is low
        if self.llm and ats_score < self.config.target_ats_score and missing:
            try:
                prompt = InternshipPrompts.format_template(
                    "ats_optimization",
                    job_keywords=", ".join(all_keywords),
                    resume_content=resume_text[:2000],
                )
                
                # Get optimization suggestions
                suggestions = await self.llm.generate(prompt)
                logger.debug(f"ATS optimization suggestions: {suggestions[:200]}...")
                
            except Exception as e:
                logger.warning(f"ATS optimization failed: {e}")
        
        keywords_result = {
            "included": included,
            "missing": missing,
        }
        
        return sections, ats_score, keywords_result
    
    def _sections_to_text(self, sections: Dict[str, Any]) -> str:
        """Convert sections to plain text for analysis."""
        parts = []
        
        if "summary" in sections:
            parts.append(sections["summary"])
        
        if "experience" in sections:
            for exp in sections["experience"]:
                parts.append(f"{exp.get('role', '')} at {exp.get('company', '')}")
                parts.extend(exp.get("bullets", []))
        
        if "projects" in sections:
            for proj in sections["projects"]:
                parts.append(proj.get("name", ""))
                parts.append(" ".join(proj.get("technologies", [])))
                parts.extend(proj.get("bullets", []))
        
        if "skills" in sections:
            for category, skills in sections["skills"].items():
                parts.extend(skills)
        
        return " ".join(parts)
    
    # =========================================================================
    # Step 5: Output Generation
    # =========================================================================
    
    async def _generate_outputs(
        self,
        resume: GeneratedResume,
        job: InternshipListing,
    ) -> GeneratedResume:
        """Generate output files in various formats."""
        # Create filename base
        safe_company = "".join(c if c.isalnum() else "_" for c in job.company)
        safe_role = "".join(c if c.isalnum() else "_" for c in job.role)
        timestamp = datetime.now().strftime("%Y%m%d")
        filename_base = f"{safe_company}_{safe_role}_{timestamp}"
        
        output_dir = Path(self.config.output_directory)
        
        # Generate each format
        for fmt in self.config.output_formats:
            try:
                if fmt == ResumeFormat.MARKDOWN:
                    path = output_dir / f"{filename_base}.md"
                    content = self._to_markdown(resume)
                    path.write_text(content, encoding="utf-8")
                    resume.full_text = content
                    
                elif fmt == ResumeFormat.PLAIN_TEXT:
                    path = output_dir / f"{filename_base}.txt"
                    content = self._to_plain_text(resume)
                    path.write_text(content, encoding="utf-8")
                    
                elif fmt == ResumeFormat.JSON:
                    path = output_dir / f"{filename_base}.json"
                    path.write_text(
                        json.dumps(resume.content, indent=2),
                        encoding="utf-8"
                    )
                    
                elif fmt == ResumeFormat.PDF:
                    path = output_dir / f"{filename_base}.pdf"
                    await self._generate_pdf(resume, path)
                    resume.pdf_path = str(path)
                    
                elif fmt == ResumeFormat.DOCX:
                    path = output_dir / f"{filename_base}.docx"
                    await self._generate_docx(resume, path)
                    resume.docx_path = str(path)
                    
                logger.debug(f"Generated {fmt.value}: {path}")
                
            except Exception as e:
                logger.error(f"Failed to generate {fmt.value}: {e}")
        
        return resume
    
    def _to_markdown(self, resume: GeneratedResume) -> str:
        """Convert resume to Markdown format."""
        lines = []
        sections = resume.content
        
        # Header
        contact = sections.get("contact", {})
        lines.append(f"# {contact.get('name', 'Resume')}")
        
        contact_parts = []
        if contact.get("email"):
            contact_parts.append(contact["email"])
        if contact.get("phone"):
            contact_parts.append(contact["phone"])
        if contact.get("linkedin"):
            contact_parts.append(f"[LinkedIn]({contact['linkedin']})")
        if contact.get("github"):
            contact_parts.append(f"[GitHub]({contact['github']})")
        
        if contact_parts:
            lines.append(" | ".join(contact_parts))
        lines.append("")
        
        # Summary
        if sections.get("summary"):
            lines.append("## Summary")
            lines.append(sections["summary"])
            lines.append("")
        
        # Education
        if sections.get("education"):
            lines.append("## Education")
            for edu in sections["education"]:
                lines.append(f"**{edu.get('institution', '')}** | {edu.get('graduation', '')}")
                lines.append(f"{edu.get('degree', '')} in {edu.get('major', '')} | GPA: {edu.get('gpa', '')}")
                if edu.get("relevant_coursework"):
                    lines.append(f"*Relevant Coursework:* {', '.join(edu['relevant_coursework'])}")
            lines.append("")
        
        # Experience
        if sections.get("experience"):
            lines.append("## Experience")
            for exp in sections["experience"]:
                lines.append(f"**{exp.get('role', '')}** | {exp.get('company', '')} | {exp.get('dates', '')}")
                for bullet in exp.get("bullets", []):
                    lines.append(f"- {bullet}")
                lines.append("")
        
        # Projects
        if sections.get("projects"):
            lines.append("## Projects")
            for proj in sections["projects"]:
                tech_str = ", ".join(proj.get("technologies", []))
                lines.append(f"**{proj.get('name', '')}** | {tech_str}")
                for bullet in proj.get("bullets", []):
                    lines.append(f"- {bullet}")
                lines.append("")
        
        # Skills
        if sections.get("skills"):
            lines.append("## Skills")
            for category, skills in sections["skills"].items():
                lines.append(f"**{category}:** {', '.join(skills)}")
            lines.append("")
        
        return "\n".join(lines)
    
    def _to_plain_text(self, resume: GeneratedResume) -> str:
        """Convert resume to plain text."""
        # Use markdown and strip formatting
        md = self._to_markdown(resume)
        
        # Simple markdown stripping
        text = md.replace("**", "").replace("*", "").replace("#", "")
        text = "\n".join(line for line in text.split("\n") if not line.startswith("["))
        
        return text
    
    async def _generate_pdf(self, resume: GeneratedResume, path: Path):
        """Generate PDF version of resume."""
        try:
            # Try using reportlab
            from reportlab.lib.pagesizes import letter
            from reportlab.pdfgen import canvas
            from reportlab.lib.units import inch
            
            c = canvas.Canvas(str(path), pagesize=letter)
            width, height = letter
            
            y = height - inch
            sections = resume.content
            
            # Name
            contact = sections.get("contact", {})
            c.setFont("Helvetica-Bold", 16)
            c.drawString(inch, y, contact.get("name", "Resume"))
            y -= 0.3 * inch
            
            # Contact info
            c.setFont("Helvetica", 10)
            contact_line = " | ".join(filter(None, [
                contact.get("email"),
                contact.get("phone"),
                contact.get("linkedin"),
            ]))
            c.drawString(inch, y, contact_line)
            y -= 0.4 * inch
            
            # Summary
            if sections.get("summary"):
                c.setFont("Helvetica-Bold", 12)
                c.drawString(inch, y, "SUMMARY")
                y -= 0.2 * inch
                c.setFont("Helvetica", 10)
                
                # Word wrap summary
                summary = sections["summary"]
                words = summary.split()
                line = ""
                for word in words:
                    if len(line + word) < 80:
                        line += word + " "
                    else:
                        c.drawString(inch, y, line.strip())
                        y -= 0.15 * inch
                        line = word + " "
                if line:
                    c.drawString(inch, y, line.strip())
                    y -= 0.3 * inch
            
            # Education
            if sections.get("education"):
                c.setFont("Helvetica-Bold", 12)
                c.drawString(inch, y, "EDUCATION")
                y -= 0.2 * inch
                
                for edu in sections["education"]:
                    c.setFont("Helvetica-Bold", 10)
                    c.drawString(inch, y, f"{edu.get('institution', '')} | {edu.get('graduation', '')}")
                    y -= 0.15 * inch
                    c.setFont("Helvetica", 10)
                    c.drawString(inch, y, f"{edu.get('degree', '')} in {edu.get('major', '')} | GPA: {edu.get('gpa', '')}")
                    y -= 0.25 * inch
            
            # Projects
            if sections.get("projects"):
                c.setFont("Helvetica-Bold", 12)
                c.drawString(inch, y, "PROJECTS")
                y -= 0.2 * inch
                
                for proj in sections["projects"]:
                    c.setFont("Helvetica-Bold", 10)
                    c.drawString(inch, y, proj.get("name", ""))
                    y -= 0.15 * inch
                    c.setFont("Helvetica", 9)
                    c.drawString(inch, y, ", ".join(proj.get("technologies", [])))
                    y -= 0.15 * inch
                    
                    for bullet in proj.get("bullets", [])[:2]:
                        c.drawString(inch + 0.2*inch, y, f"• {bullet[:80]}...")
                        y -= 0.15 * inch
                    y -= 0.1 * inch
            
            # Skills
            if sections.get("skills"):
                c.setFont("Helvetica-Bold", 12)
                c.drawString(inch, y, "SKILLS")
                y -= 0.2 * inch
                c.setFont("Helvetica", 10)
                
                for category, skills in sections["skills"].items():
                    c.drawString(inch, y, f"{category}: {', '.join(skills)}")
                    y -= 0.15 * inch
            
            c.save()
            logger.info(f"PDF generated: {path}")
            
        except ImportError:
            logger.warning("reportlab not installed - PDF generation skipped")
        except Exception as e:
            logger.error(f"PDF generation failed: {e}")
    
    async def _generate_docx(self, resume: GeneratedResume, path: Path):
        """Generate DOCX version of resume."""
        try:
            from docx import Document
            from docx.shared import Pt, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            doc = Document()
            sections = resume.content
            
            # Name
            contact = sections.get("contact", {})
            name_para = doc.add_paragraph()
            name_run = name_para.add_run(contact.get("name", "Resume"))
            name_run.bold = True
            name_run.font.size = Pt(18)
            name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Contact
            contact_para = doc.add_paragraph()
            contact_parts = [
                contact.get("email", ""),
                contact.get("phone", ""),
                contact.get("linkedin", ""),
            ]
            contact_para.add_run(" | ".join(filter(None, contact_parts)))
            contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Summary
            if sections.get("summary"):
                doc.add_heading("Summary", level=1)
                doc.add_paragraph(sections["summary"])
            
            # Education
            if sections.get("education"):
                doc.add_heading("Education", level=1)
                for edu in sections["education"]:
                    p = doc.add_paragraph()
                    p.add_run(f"{edu.get('institution', '')}").bold = True
                    p.add_run(f" | {edu.get('graduation', '')}")
                    doc.add_paragraph(
                        f"{edu.get('degree', '')} in {edu.get('major', '')} | GPA: {edu.get('gpa', '')}"
                    )
            
            # Projects
            if sections.get("projects"):
                doc.add_heading("Projects", level=1)
                for proj in sections["projects"]:
                    p = doc.add_paragraph()
                    p.add_run(proj.get("name", "")).bold = True
                    p.add_run(f" | {', '.join(proj.get('technologies', []))}")
                    
                    for bullet in proj.get("bullets", []):
                        doc.add_paragraph(bullet, style="List Bullet")
            
            # Skills
            if sections.get("skills"):
                doc.add_heading("Skills", level=1)
                for category, skills in sections["skills"].items():
                    p = doc.add_paragraph()
                    p.add_run(f"{category}: ").bold = True
                    p.add_run(", ".join(skills))
            
            doc.save(str(path))
            logger.info(f"DOCX generated: {path}")
            
        except ImportError:
            logger.warning("python-docx not installed - DOCX generation skipped")
        except Exception as e:
            logger.error(f"DOCX generation failed: {e}")
    
    # =========================================================================
    # Utility Methods
    # =========================================================================
    
    def get_generation_summary(self, resume: GeneratedResume) -> str:
        """Get a summary of the generated resume."""
        lines = [
            f"📝 **Resume Generated for {resume.company} - {resume.role}**",
            "",
            f"**ATS Score:** {resume.ats_score:.0f}%",
            "",
            "**Keywords Included:**",
            f"  {', '.join(resume.keywords_included[:10])}",
            "",
        ]
        
        if resume.keywords_missing:
            lines.extend([
                "**Keywords Missing:**",
                f"  {', '.join(resume.keywords_missing[:5])}",
                "",
            ])
        
        lines.extend([
            "**Files Generated:**",
        ])
        
        if resume.pdf_path:
            lines.append(f"  📄 PDF: {resume.pdf_path}")
        if resume.docx_path:
            lines.append(f"  📄 DOCX: {resume.docx_path}")
        
        return "\n".join(lines)
