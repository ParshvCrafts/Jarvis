"""
Document/RAG Service for JARVIS - Phase 7

Provides document ingestion and retrieval-augmented generation (RAG).

Features:
- PDF document parsing
- Text file ingestion
- Web page content extraction
- Semantic chunking
- Vector search with ChromaDB
- Question answering with citations

Dependencies:
- chromadb: Vector database
- pypdf2 or pdfplumber: PDF parsing
- sentence-transformers: Embeddings (optional, uses ChromaDB default)
"""

from __future__ import annotations

import asyncio
import hashlib
import os
import re
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple
from enum import Enum

from loguru import logger

# ChromaDB (optional)
try:
    import chromadb
    from chromadb.config import Settings
    CHROMADB_AVAILABLE = True
except ImportError:
    CHROMADB_AVAILABLE = False
    logger.debug("chromadb not installed. Install with: pip install chromadb")

# PDF parsing (optional)
try:
    from pypdf import PdfReader
    PDF_AVAILABLE = True
except ImportError:
    try:
        from PyPDF2 import PdfReader
        PDF_AVAILABLE = True
    except ImportError:
        PDF_AVAILABLE = False
        logger.debug("PDF parsing not available. Install with: pip install pypdf")

# Word document parsing (optional)
try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.debug("DOCX parsing not available. Install with: pip install python-docx")


class DocumentType(Enum):
    """Supported document types."""
    PDF = "pdf"
    TEXT = "txt"
    MARKDOWN = "md"
    DOCX = "docx"
    HTML = "html"
    UNKNOWN = "unknown"


@dataclass
class DocumentChunk:
    """A chunk of document content."""
    id: str
    content: str
    document_id: str
    document_name: str
    chunk_index: int
    metadata: Dict[str, Any] = field(default_factory=dict)
    
    @property
    def preview(self) -> str:
        """Get a preview of the chunk content."""
        return self.content[:200] + "..." if len(self.content) > 200 else self.content


@dataclass
class Document:
    """Represents an ingested document."""
    id: str
    name: str
    source: str  # File path or URL
    doc_type: DocumentType
    content: str
    chunks: List[DocumentChunk] = field(default_factory=list)
    metadata: Dict[str, Any] = field(default_factory=dict)
    ingested_at: datetime = field(default_factory=datetime.now)
    
    @property
    def chunk_count(self) -> int:
        return len(self.chunks)
    
    def format_info(self) -> str:
        """Format document info for display."""
        return (
            f"📄 **{self.name}**\n"
            f"   Type: {self.doc_type.value}\n"
            f"   Chunks: {self.chunk_count}\n"
            f"   Ingested: {self.ingested_at.strftime('%Y-%m-%d %H:%M')}"
        )


@dataclass
class SearchResult:
    """A search result from the vector store."""
    chunk: DocumentChunk
    score: float
    
    def format_citation(self) -> str:
        """Format as a citation."""
        return f"[{self.chunk.document_name}, chunk {self.chunk.chunk_index + 1}]"


class TextChunker:
    """
    Splits text into semantic chunks for vector storage.
    
    Uses sentence-aware chunking with overlap for context continuity.
    """
    
    def __init__(
        self,
        chunk_size: int = 500,
        chunk_overlap: int = 50,
        min_chunk_size: int = 100,
    ):
        """
        Initialize chunker.
        
        Args:
            chunk_size: Target chunk size in characters
            chunk_overlap: Overlap between chunks
            min_chunk_size: Minimum chunk size
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.min_chunk_size = min_chunk_size
    
    def chunk_text(self, text: str) -> List[str]:
        """
        Split text into chunks.
        
        Args:
            text: Text to chunk
            
        Returns:
            List of text chunks
        """
        # Clean text
        text = re.sub(r'\s+', ' ', text).strip()
        
        if len(text) <= self.chunk_size:
            return [text] if len(text) >= self.min_chunk_size else []
        
        # Split into sentences
        sentences = re.split(r'(?<=[.!?])\s+', text)
        
        chunks = []
        current_chunk = []
        current_length = 0
        
        for sentence in sentences:
            sentence_length = len(sentence)
            
            if current_length + sentence_length > self.chunk_size and current_chunk:
                # Save current chunk
                chunk_text = ' '.join(current_chunk)
                if len(chunk_text) >= self.min_chunk_size:
                    chunks.append(chunk_text)
                
                # Start new chunk with overlap
                overlap_text = ' '.join(current_chunk[-2:]) if len(current_chunk) > 1 else ''
                current_chunk = [overlap_text] if overlap_text else []
                current_length = len(overlap_text)
            
            current_chunk.append(sentence)
            current_length += sentence_length + 1
        
        # Add final chunk
        if current_chunk:
            chunk_text = ' '.join(current_chunk)
            if len(chunk_text) >= self.min_chunk_size:
                chunks.append(chunk_text)
        
        return chunks


class DocumentParser:
    """Parses various document formats into text."""
    
    @staticmethod
    def detect_type(path: str) -> DocumentType:
        """Detect document type from file extension."""
        ext = Path(path).suffix.lower()
        type_map = {
            ".pdf": DocumentType.PDF,
            ".txt": DocumentType.TEXT,
            ".md": DocumentType.MARKDOWN,
            ".docx": DocumentType.DOCX,
            ".html": DocumentType.HTML,
            ".htm": DocumentType.HTML,
        }
        return type_map.get(ext, DocumentType.UNKNOWN)
    
    @staticmethod
    def parse_pdf(path: str) -> str:
        """Parse PDF file to text."""
        if not PDF_AVAILABLE:
            raise ImportError("PDF parsing not available. Install pypdf.")
        
        reader = PdfReader(path)
        text_parts = []
        
        for page in reader.pages:
            text = page.extract_text()
            if text:
                text_parts.append(text)
        
        return "\n\n".join(text_parts)
    
    @staticmethod
    def parse_docx(path: str) -> str:
        """Parse DOCX file to text."""
        if not DOCX_AVAILABLE:
            raise ImportError("DOCX parsing not available. Install python-docx.")
        
        doc = DocxDocument(path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return "\n\n".join(paragraphs)
    
    @staticmethod
    def parse_text(path: str) -> str:
        """Parse text file."""
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    
    @staticmethod
    def parse_html(content: str) -> str:
        """Parse HTML to plain text."""
        # Simple HTML tag removal
        text = re.sub(r'<script[^>]*>.*?</script>', '', content, flags=re.DOTALL)
        text = re.sub(r'<style[^>]*>.*?</style>', '', text, flags=re.DOTALL)
        text = re.sub(r'<[^>]+>', ' ', text)
        text = re.sub(r'\s+', ' ', text)
        return text.strip()
    
    @classmethod
    def parse_file(cls, path: str) -> Tuple[str, DocumentType]:
        """
        Parse a file to text.
        
        Args:
            path: File path
            
        Returns:
            Tuple of (text content, document type)
        """
        doc_type = cls.detect_type(path)
        
        if doc_type == DocumentType.PDF:
            content = cls.parse_pdf(path)
        elif doc_type == DocumentType.DOCX:
            content = cls.parse_docx(path)
        elif doc_type in [DocumentType.TEXT, DocumentType.MARKDOWN]:
            content = cls.parse_text(path)
        elif doc_type == DocumentType.HTML:
            raw = cls.parse_text(path)
            content = cls.parse_html(raw)
        else:
            # Try as text
            content = cls.parse_text(path)
            doc_type = DocumentType.TEXT
        
        return content, doc_type


class DocumentService:
    """
    Document service for RAG (Retrieval-Augmented Generation).
    
    Provides document ingestion, vector storage, and semantic search
    for question answering with citations.
    """
    
    def __init__(
        self,
        persist_directory: str = "data/documents",
        collection_name: str = "jarvis_documents",
        chunk_size: int = 500,
        chunk_overlap: int = 50,
    ):
        """
        Initialize document service.
        
        Args:
            persist_directory: Directory for vector store persistence
            collection_name: ChromaDB collection name
            chunk_size: Chunk size for text splitting
            chunk_overlap: Overlap between chunks
        """
        self.persist_directory = Path(persist_directory)
        self.collection_name = collection_name
        self.chunker = TextChunker(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
        self.parser = DocumentParser()
        
        self._client = None
        self._collection = None
        self._documents: Dict[str, Document] = {}
    
    def _get_client(self):
        """Get or create ChromaDB client."""
        if not CHROMADB_AVAILABLE:
            raise ImportError("ChromaDB not available. Install with: pip install chromadb")
        
        if self._client is None:
            self.persist_directory.mkdir(parents=True, exist_ok=True)
            self._client = chromadb.PersistentClient(
                path=str(self.persist_directory),
                settings=Settings(anonymized_telemetry=False),
            )
        
        return self._client
    
    def _get_collection(self):
        """Get or create ChromaDB collection."""
        if self._collection is None:
            client = self._get_client()
            self._collection = client.get_or_create_collection(
                name=self.collection_name,
                metadata={"hnsw:space": "cosine"},
            )
        return self._collection
    
    def is_available(self) -> bool:
        """Check if document service is available."""
        return CHROMADB_AVAILABLE
    
    def _generate_id(self, content: str) -> str:
        """Generate unique ID from content."""
        return hashlib.md5(content.encode()).hexdigest()[:12]
    
    async def ingest_file(
        self,
        file_path: str,
        metadata: Optional[Dict[str, Any]] = None,
    ) -> Optional[Document]:
        """
        Ingest a document file.
        
        Args:
            file_path: Path to the document
            metadata: Optional metadata to attach
            
        Returns:
            Ingested Document or None on failure
        """
        try:
            path = Path(file_path)
            if not path.exists():
                logger.error(f"File not found: {file_path}")
                return None
            
            # Parse document
            content, doc_type = self.parser.parse_file(str(path))
            
            if not content.strip():
                logger.warning(f"No content extracted from {file_path}")
                return None
            
            # Create document
            doc_id = self._generate_id(str(path) + content[:100])
            document = Document(
                id=doc_id,
                name=path.name,
                source=str(path),
                doc_type=doc_type,
                content=content,
                metadata=metadata or {},
            )
            
            # Chunk content
            chunk_texts = self.chunker.chunk_text(content)
            
            # Create chunks
            chunks = []
            for i, chunk_text in enumerate(chunk_texts):
                chunk = DocumentChunk(
                    id=f"{doc_id}_{i}",
                    content=chunk_text,
                    document_id=doc_id,
                    document_name=document.name,
                    chunk_index=i,
                    metadata={
                        "source": str(path),
                        "doc_type": doc_type.value,
                        **(metadata or {}),
                    },
                )
                chunks.append(chunk)
            
            document.chunks = chunks
            
            # Add to vector store
            collection = self._get_collection()
            
            # Run in executor to avoid blocking
            loop = asyncio.get_event_loop()
            await loop.run_in_executor(
                None,
                lambda: collection.add(
                    ids=[c.id for c in chunks],
                    documents=[c.content for c in chunks],
                    metadatas=[c.metadata for c in chunks],
                )
            )
            
            # Store document reference
            self._documents[doc_id] = document
            
            logger.info(f"Ingested document: {document.name} ({len(chunks)} chunks)")
            return document
            
        except Exception as e:
            logger.error(f"Failed to ingest {file_path}: {e}")
            return None
    
    async def ingest_text(
        self,
        text: str,
        name: str,
        metadata: Optional[Dict[str, Any]] = None,
    ) -> Optional[Document]:
        """
        Ingest raw text content.
        
        Args:
            text: Text content
            name: Document name
            metadata: Optional metadata
            
        Returns:
            Ingested Document or None on failure
        """
        try:
            doc_id = self._generate_id(name + text[:100])
            document = Document(
                id=doc_id,
                name=name,
                source="text",
                doc_type=DocumentType.TEXT,
                content=text,
                metadata=metadata or {},
            )
            
            # Chunk content
            chunk_texts = self.chunker.chunk_text(text)
            
            chunks = []
            for i, chunk_text in enumerate(chunk_texts):
                chunk = DocumentChunk(
                    id=f"{doc_id}_{i}",
                    content=chunk_text,
                    document_id=doc_id,
                    document_name=name,
                    chunk_index=i,
                    metadata=metadata or {},
                )
                chunks.append(chunk)
            
            document.chunks = chunks
            
            # Add to vector store
            collection = self._get_collection()
            loop = asyncio.get_event_loop()
            await loop.run_in_executor(
                None,
                lambda: collection.add(
                    ids=[c.id for c in chunks],
                    documents=[c.content for c in chunks],
                    metadatas=[c.metadata for c in chunks],
                )
            )
            
            self._documents[doc_id] = document
            logger.info(f"Ingested text: {name} ({len(chunks)} chunks)")
            return document
            
        except Exception as e:
            logger.error(f"Failed to ingest text: {e}")
            return None
    
    async def search(
        self,
        query: str,
        top_k: int = 5,
        filter_metadata: Optional[Dict[str, Any]] = None,
    ) -> List[SearchResult]:
        """
        Search for relevant document chunks.
        
        Args:
            query: Search query
            top_k: Number of results to return
            filter_metadata: Optional metadata filter
            
        Returns:
            List of SearchResult objects
        """
        try:
            collection = self._get_collection()
            
            loop = asyncio.get_event_loop()
            results = await loop.run_in_executor(
                None,
                lambda: collection.query(
                    query_texts=[query],
                    n_results=top_k,
                    where=filter_metadata,
                )
            )
            
            search_results = []
            
            if results and results["ids"] and results["ids"][0]:
                for i, chunk_id in enumerate(results["ids"][0]):
                    # Find the chunk
                    content = results["documents"][0][i] if results["documents"] else ""
                    metadata = results["metadatas"][0][i] if results["metadatas"] else {}
                    distance = results["distances"][0][i] if results["distances"] else 0
                    
                    # Convert distance to similarity score
                    score = 1 - distance  # For cosine distance
                    
                    chunk = DocumentChunk(
                        id=chunk_id,
                        content=content,
                        document_id=metadata.get("document_id", ""),
                        document_name=metadata.get("source", "Unknown").split("/")[-1],
                        chunk_index=int(chunk_id.split("_")[-1]) if "_" in chunk_id else 0,
                        metadata=metadata,
                    )
                    
                    search_results.append(SearchResult(chunk=chunk, score=score))
            
            return search_results
            
        except Exception as e:
            logger.error(f"Search failed: {e}")
            return []
    
    async def query_documents(
        self,
        question: str,
        top_k: int = 5,
    ) -> Tuple[str, List[SearchResult]]:
        """
        Query documents and return relevant context.
        
        Args:
            question: User question
            top_k: Number of chunks to retrieve
            
        Returns:
            Tuple of (context string, search results)
        """
        results = await self.search(question, top_k=top_k)
        
        if not results:
            return "", []
        
        # Build context from results
        context_parts = []
        for i, result in enumerate(results, 1):
            context_parts.append(
                f"[Source {i}: {result.chunk.document_name}]\n{result.chunk.content}"
            )
        
        context = "\n\n---\n\n".join(context_parts)
        return context, results
    
    def list_documents(self) -> List[Document]:
        """List all ingested documents."""
        return list(self._documents.values())
    
    async def delete_document(self, doc_id: str) -> bool:
        """
        Delete a document and its chunks.
        
        Args:
            doc_id: Document ID
            
        Returns:
            True if deleted successfully
        """
        try:
            if doc_id not in self._documents:
                return False
            
            document = self._documents[doc_id]
            chunk_ids = [c.id for c in document.chunks]
            
            # Delete from vector store
            collection = self._get_collection()
            loop = asyncio.get_event_loop()
            await loop.run_in_executor(
                None,
                lambda: collection.delete(ids=chunk_ids)
            )
            
            # Remove from cache
            del self._documents[doc_id]
            
            logger.info(f"Deleted document: {document.name}")
            return True
            
        except Exception as e:
            logger.error(f"Failed to delete document: {e}")
            return False


# Singleton instance
_document_service: Optional[DocumentService] = None


def get_document_service(
    persist_directory: str = "data/documents",
) -> DocumentService:
    """Get or create document service singleton."""
    global _document_service
    if _document_service is None:
        _document_service = DocumentService(persist_directory=persist_directory)
    return _document_service


# =============================================================================
# Tool Functions for Agent Integration
# =============================================================================

async def ingest_document(file_path: str) -> str:
    """
    Ingest a document for RAG queries.
    
    Args:
        file_path: Path to the document (PDF, TXT, DOCX, MD)
        
    Returns:
        Confirmation message
    """
    try:
        service = get_document_service()
        
        if not service.is_available():
            return "Document service not available. Install chromadb: pip install chromadb"
        
        document = await service.ingest_file(file_path)
        
        if document:
            return (
                f"✅ Ingested document: **{document.name}**\n"
                f"   Type: {document.doc_type.value}\n"
                f"   Chunks: {document.chunk_count}\n"
                f"   Ready for questions!"
            )
        else:
            return f"Failed to ingest document: {file_path}"
        
    except Exception as e:
        logger.error(f"Failed to ingest document: {e}")
        return f"Failed to ingest document: {e}"


async def query_documents(question: str) -> str:
    """
    Ask a question about ingested documents.
    
    Args:
        question: Question to ask
        
    Returns:
        Answer with citations
    """
    try:
        service = get_document_service()
        
        if not service.is_available():
            return "Document service not available."
        
        context, results = await service.query_documents(question, top_k=5)
        
        if not results:
            return "No relevant information found in your documents. Try ingesting more documents or rephrasing your question."
        
        # Format response with context and citations
        lines = ["**Relevant Information Found:**\n"]
        
        for i, result in enumerate(results, 1):
            lines.append(f"**Source {i}:** {result.chunk.document_name}")
            lines.append(f"_{result.chunk.preview}_")
            lines.append("")
        
        lines.append("---")
        lines.append(f"*Based on {len(results)} document chunks. Relevance scores: {', '.join([f'{r.score:.2f}' for r in results])}*")
        
        return "\n".join(lines)
        
    except Exception as e:
        logger.error(f"Failed to query documents: {e}")
        return f"Failed to query documents: {e}"


async def list_ingested_documents() -> str:
    """
    List all ingested documents.
    
    Returns:
        List of documents
    """
    try:
        service = get_document_service()
        
        if not service.is_available():
            return "Document service not available."
        
        documents = service.list_documents()
        
        if not documents:
            return "No documents ingested yet. Use 'ingest document' to add documents."
        
        lines = [f"**Ingested Documents** ({len(documents)} total)\n"]
        for doc in documents:
            lines.append(doc.format_info())
            lines.append("")
        
        return "\n".join(lines)
        
    except Exception as e:
        logger.error(f"Failed to list documents: {e}")
        return f"Failed to list documents: {e}"


async def delete_document(document_name: str) -> str:
    """
    Delete an ingested document.
    
    Args:
        document_name: Name of the document to delete
        
    Returns:
        Confirmation message
    """
    try:
        service = get_document_service()
        
        if not service.is_available():
            return "Document service not available."
        
        # Find document by name
        documents = service.list_documents()
        matching = [d for d in documents if document_name.lower() in d.name.lower()]
        
        if not matching:
            return f"Document '{document_name}' not found."
        
        if len(matching) > 1:
            names = [d.name for d in matching]
            return f"Multiple documents match '{document_name}': {', '.join(names)}. Please be more specific."
        
        doc = matching[0]
        success = await service.delete_document(doc.id)
        
        if success:
            return f"✅ Deleted document: {doc.name}"
        else:
            return f"Failed to delete document: {doc.name}"
        
    except Exception as e:
        logger.error(f"Failed to delete document: {e}")
        return f"Failed to delete document: {e}"


# Document tool definitions for agent system
DOCUMENT_TOOLS = [
    {
        "name": "ingest_document",
        "description": "Ingest a document (PDF, TXT, DOCX, MD) for question answering. Use this when the user wants to add a document to their knowledge base.",
        "parameters": {
            "type": "object",
            "properties": {
                "file_path": {
                    "type": "string",
                    "description": "Full path to the document file"
                }
            },
            "required": ["file_path"]
        },
        "function": ingest_document,
    },
    {
        "name": "query_documents",
        "description": "Ask a question about ingested documents. Use this when the user asks about their documents or uploaded files.",
        "parameters": {
            "type": "object",
            "properties": {
                "question": {
                    "type": "string",
                    "description": "Question to ask about the documents"
                }
            },
            "required": ["question"]
        },
        "function": query_documents,
    },
    {
        "name": "list_ingested_documents",
        "description": "List all documents that have been ingested for RAG queries.",
        "parameters": {
            "type": "object",
            "properties": {},
            "required": []
        },
        "function": list_ingested_documents,
    },
    {
        "name": "delete_document",
        "description": "Delete an ingested document from the knowledge base.",
        "parameters": {
            "type": "object",
            "properties": {
                "document_name": {
                    "type": "string",
                    "description": "Name of the document to delete"
                }
            },
            "required": ["document_name"]
        },
        "function": delete_document,
    },
]
